import streamlit as st
import os
from git import Repo, Actor
from datetime import datetime
import shutil
import stat
import re
from dotenv import load_dotenv

# --- Docx and LangChain Imports ---
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument

# --- Load .env for local development ---
load_dotenv()

# --- File Paths and URLs ---
REPO_URL = "https://github.com/BasheerKalanchera/Halal-Investment-App-Data-Converter.git"
FAISS_DIR = "faiss_index"
LOCAL_CLONE_PATH = "./temp_repo"

# --- GitHub Authentication ---
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN") or st.secrets.get("GITHUB_TOKEN")
if not GITHUB_TOKEN:
    st.error("GitHub Personal Access Token not found...")
    st.stop()

# --- Helper Functions ---
def remove_readonly(func, path, excinfo):
    """Error handler for shutil.rmtree."""
    os.chmod(path, stat.S_IWRITE)
    func(path)

# --- Embedding model ---
@st.cache_resource
def load_embedding_model():
    """Loads the HuggingFace embedding model."""
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={"device": "cpu"}
    )

# --- Document Parser ---
def parse_docx_for_qas(file_path):
    try:
        doc = DocxDocument(file_path)
        qa_pairs = []
        q_text = None
        a_text_parts = []

        for block in doc.element.body:
            if isinstance(block, CT_P):
                p = Paragraph(block, doc)
                p_text = p.text.strip()
                if p_text.lower() == 'question':
                    if q_text and a_text_parts:
                        answer = "\n".join(a_text_parts).strip()
                        qa_pairs.append(LangchainDocument(page_content=answer, metadata={"question": q_text}))
                    q_text = None
                    a_text_parts = []
                elif q_text is None and p_text:
                    q_text = p_text
                elif q_text and p_text:
                    a_text_parts.append(p_text)
            elif isinstance(block, CT_Tbl):
                if q_text:
                    t = Table(block, doc)
                    table_data = [",".join([cell.text.replace('\n', ' ').strip() for cell in row.cells]) for row in t.rows]
                    a_text_parts.append("\n".join(table_data))
        
        if q_text and a_text_parts:
            answer = "\n".join(a_text_parts).strip()
            qa_pairs.append(LangchainDocument(page_content=answer, metadata={"question": q_text}))

        return qa_pairs
    except Exception as e:
        st.error(f"Error parsing .docx file: {e}")
        return []

# --- Main app logic ---
def run_converter_app():
    st.title("📚 Halal Gold Knowledge Base Converter")
    st.info("Upload a .docx file to update the app's knowledge base on GitHub.")
    uploaded_file = st.file_uploader("Choose a .docx file", type=["docx"])

    if uploaded_file and st.button("Update Knowledge Base"):
        with st.spinner("Processing..."):
            repo_path = LOCAL_CLONE_PATH
            repo_with_token = REPO_URL.replace("https://", f"https://oauth2:{GITHUB_TOKEN}@")

            st.write("Syncing with knowledge base repository from GitHub...")
            if os.path.exists(repo_path):
                shutil.rmtree(repo_path, onerror=remove_readonly)
            
            repo = Repo.clone_from(repo_with_token, repo_path)
            full_faiss_path = os.path.join(repo_path, FAISS_DIR)
            embeddings = load_embedding_model()
            
            st.write("Loading existing knowledge base...")
            existing_docs_map = {}
            if os.path.exists(full_faiss_path):
                try:
                    vectorstore = FAISS.load_local(full_faiss_path, embeddings, allow_dangerous_deserialization=True)
                    for doc in vectorstore.docstore._dict.values():
                        if 'question' in doc.metadata:
                            existing_docs_map[doc.metadata['question']] = doc
                    st.success(f"Successfully loaded {len(existing_docs_map)} existing documents.")
                except Exception as e:
                    st.warning(f"Could not load existing index, will create a new one. Error: {e}")
            else:
                st.info("No existing knowledge base found.")

            temp_docx_path = f"temp_{uploaded_file.name}"
            with open(temp_docx_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            new_docs_from_file = parse_docx_for_qas(temp_docx_path)
            os.remove(temp_docx_path)
            st.success(f"Parsed {len(new_docs_from_file)} Q&A pairs from the .docx file.")

            updates = 0
            additions = 0
            for new_doc in new_docs_from_file:
                question = new_doc.metadata["question"]
                new_content_normalized = re.sub(r'\s+', ' ', new_doc.page_content).strip()

                if question in existing_docs_map:
                    existing_content_normalized = re.sub(r'\s+', ' ', existing_docs_map[question].page_content).strip()
                    if new_content_normalized != existing_content_normalized:
                        updates += 1
                        existing_docs_map[question] = new_doc
                else:
                    additions += 1
                    existing_docs_map[question] = new_doc

            if updates == 0 and additions == 0:
                st.info("The knowledge base is already up-to-date.")
                st.stop()
            
            st.write(f"Processing: {additions} new document(s) and {updates} updated document(s).")
            
            # --- FINAL FIX: Create a clean list of documents before rebuilding ---
            st.write("Purifying final document list...")
            clean_final_documents = []
            for doc in existing_docs_map.values():
                clean_final_documents.append(
                    LangchainDocument(
                        page_content=doc.page_content,
                        metadata={"question": doc.metadata["question"]}
                    )
                )

            st.write("Rebuilding knowledge base...")
            if not clean_final_documents:
                 st.warning("No documents to build. Stopping.")
                 st.stop()
            
            new_vectorstore = FAISS.from_documents(clean_final_documents, embeddings)
            
            try:
                st.write("Saving updated index and pushing to GitHub...")
                new_vectorstore.save_local(full_faiss_path)
                
                repo.index.add([os.path.join(FAISS_DIR, "*")])
                if repo.is_dirty(index=True, working_tree=False, untracked_files=True):
                    author = Actor("Converter Bot", "converter-bot@example.com")
                    repo.index.commit("Automated knowledge base update", author=author)
                    repo.remote(name='origin').push()
                    st.success("Knowledge base successfully updated and pushed to GitHub!")
                else:
                    st.info("No net changes to commit. Repo is clean.")
            except Exception as e:
                st.error(f"Failed to push to GitHub: {e}")

if __name__ == "__main__":
    run_converter_app()